from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches
import os
import re
import requests
from datetime import datetime

UPLOAD_FOLDER = "uploads"
DOCX_FOLDER = "generated_docs"


def download_image(url, save_path):
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
        }
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            with open(save_path, "wb") as f:
                f.write(response.content)
            return True
        else:
            print(f"Erro ao baixar imagem: {url} - Status: {response.status_code}")
            return False
    except Exception as e:
        print(f"Erro ao baixar imagem: {url} - {str(e)}")
        return False


def generate_docx_from_result(task_id, task_result):
    os.makedirs(DOCX_FOLDER, exist_ok=True)
    temp_image_folder = os.path.join(DOCX_FOLDER, f"images_{task_id}")
    os.makedirs(temp_image_folder, exist_ok=True)

    doc = Document()

    # ✅ CAPA
    doc.add_heading('Relatório de Atividades Diagnósticas', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Task ID: {task_id}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Data de geração: {datetime.now().strftime("%d/%m/%Y %H:%M")}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    trilha_count = 1
    for bloco in task_result:
        texto = bloco.get("texto", "").strip()
        opcoes = bloco.get("opcoes", []) or []
        imagens = bloco.get("imagens_url", []) or []

        # 🔹 Título da trilha
        doc.add_heading(f'Trilha {trilha_count}', level=1)
        trilha_count += 1

        # 🔹 Enunciado
        if texto:
            para = doc.add_paragraph(texto)
            para.paragraph_format.space_after = Pt(10)

        # 🔹 Imagens (após o texto, antes das opções)
        for url in imagens:
            if not url.startswith("http"):
                continue
            filename = os.path.join(temp_image_folder, os.path.basename(url.split("?")[0]))
            if download_image(url, filename):
                try:
                    doc.add_paragraph()
                    doc.add_picture(filename, width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption = doc.add_paragraph("[Imagem ilustrativa da atividade]")
                    caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    caption.paragraph_format.space_after = Pt(10)
                except Exception as e:
                    doc.add_paragraph(f"[Erro ao inserir imagem: {str(e)}]")

        # 🔹 Opções (bullets ou linhas de resposta)
        if opcoes:
            for opcao in opcoes:
                doc.add_paragraph(opcao, style='List Bullet')
        else:
            for _ in range(3):
                doc.add_paragraph("__" * 25)

        # 🔹 Separador visual
        doc.add_paragraph("\n" + ("-" * 100) + "\n")

    output_path = os.path.join(DOCX_FOLDER, f"{task_id}.docx")
    doc.save(output_path)
    return output_path
